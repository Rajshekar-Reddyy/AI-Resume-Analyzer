from pathlib import Path


class TextExtractionError(ValueError):
    """Raised when a resume or job description file cannot be parsed."""


def extract_text_from_upload(uploaded_file) -> str:
    extension = Path(uploaded_file.name).suffix.lower()
    uploaded_file.seek(0)

    if extension == ".pdf":
        text = _extract_pdf_text(uploaded_file)
    elif extension == ".docx":
        text = _extract_docx_text(uploaded_file)
    elif extension == ".txt":
        text = uploaded_file.read().decode("utf-8", errors="ignore")
    else:
        raise TextExtractionError("Unsupported file type.")

    uploaded_file.seek(0)
    cleaned = normalize_whitespace(text)
    if not cleaned:
        raise TextExtractionError(
            "No readable text was found. Try a text-based PDF or DOCX file."
        )
    return cleaned


def normalize_whitespace(text: str) -> str:
    return " ".join((text or "").replace("\x00", " ").split())


def _extract_pdf_text(uploaded_file) -> str:
    try:
        import pdfplumber

        with pdfplumber.open(uploaded_file) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages)
    except Exception:
        uploaded_file.seek(0)
        try:
            import fitz

            data = uploaded_file.read()
            with fitz.open(stream=data, filetype="pdf") as document:
                return "\n".join(page.get_text("text") for page in document)
        except Exception as exc:
            raise TextExtractionError(
                "Could not extract text from this PDF. Try another PDF or DOCX file."
            ) from exc


def _extract_docx_text(uploaded_file) -> str:
    try:
        from docx import Document

        document = Document(uploaded_file)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                table_cells.extend(cell.text for cell in row.cells)
        return "\n".join(paragraphs + table_cells)
    except Exception as exc:
        raise TextExtractionError(
            "Could not extract text from this DOCX file. Try saving it again from Word."
        ) from exc
